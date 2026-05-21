from docx import Document
import pandas as pd
import sys
import os
from openpyxl.styles import Alignment



def map_table_cells(table, start_row, start_col,new_row_flag, cell_map):
    """
    递归遍历表格（包括嵌套表格），将每个单元格的文本映射到全局的 (行, 列) 坐标上。
    """
    for row_idx, row in enumerate(table.rows):
        current_global_row = start_row + row_idx
        current_col_offset = 0
        global_col = start_col + current_col_offset
        row_adjust = True

        for cell in row.cells:
            # 计算当前单元格的实际列位置（考虑前面的合并单元格占位）
            #while (current_global_row, start_col + current_col_offset) in cell_map:
            #    current_col_offset += 1

            #global_col = start_col + current_col_offset

            # 如果单元格内包含嵌套表格，进行递归处理
            if cell.tables:
                for inner_table in cell.tables:
                    # 嵌套表格通常从下一行、当前列开始布局（也可以根据需求调整为 current_global_row, global_col）
                    map_table_cells(inner_table, current_global_row, global_col, False,cell_map)
                    global_col =global_col+1
                    #a =1
            else:
                # 如果是普通单元格，将其文本存入坐标映射字典
                if new_row_flag and row_adjust:
                    current_global_row = current_global_row*2
                    row_adjust = False
                if new_row_flag and global_col==5:
                        global_col = global_col+30

                # 只有当该位置还没有内容时才写入（防止跨行/跨列单元格的重复填充冲突）
                if (current_global_row, global_col) not in cell_map:
                    cell_map[(current_global_row, global_col)] = cell.text.strip()
                    global_col +=1

            # 考虑当前单元格本身的水平合并（grid_span）
            #grid_span = cell.grid_span if hasattr(cell, 'grid_span') else 1
            #current_col_offset += grid_span


def extract_data_from_header_footer_tables(doc_path):
    doc = Document(doc_path)
    extracted_data = []

    # 遍历文档中的每一个“节”(Section)
    for section in doc.sections:
        # 检查页眉
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    # 调用解析函数处理单元格文本
                    extracted_data.append(text)

                    target_cell = cell
                    inner_table = target_cell.tables
                    if inner_table:
                        for inner_table in cell.tables:
                            for row in inner_table.rows:
                                for cell in row.cells:
                                    text = cell.text.strip()
                                    # 调用解析函数处理单元格文本
                                    extracted_data.append(text)



        # 检查页脚
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    extracted_data.append(text)

    return extracted_data



# 获取 exe 所在的文件夹目录
exe_dir = os.path.dirname(sys.executable)

# 让使用者在弹出的黑框里输入具体的文件名
# 1. 获取用户输入并去除首尾可能多打的空格
input_name = input("请输入你要处理的 Word 文件名（包含后缀，如 报告.docx）：").strip()
# 2. 智能补充后缀：如果用户没输 .docx，自动帮他加上
if not input_name.lower().endswith('.docx'):
    input_name += '.docx'

target_name = input_name
word_file_path = os.path.join(exe_dir, target_name.strip()) # strip() 去除可能多输入的空格

if os.path.exists(word_file_path):
    doc = Document(word_file_path)
    print(f"成功读取 {target_name}，正在处理中...")

    # 使用字典来保存整个 Excel 画布上的 (行, 列) -> 内容 的映射
    all_cells = {}
    current_excel_row = 0

    # 遍历文档中的所有顶层表格
    for table in doc.tables:
        # 提取当前表格，填入 all_cells 映射中
        map_table_cells(table, current_excel_row, 0, True,all_cells)

        # 更新下一个表格在 Excel 中应该开始的行号（留出一定的间隔或紧挨着）
        if all_cells:
            max_row = max(r for r, c in all_cells.keys())
            current_excel_row = max_row + 2  # 每个大表格之间空一行

    # 将坐标映射转换为二维列表，以便生成 DataFrame
    if all_cells:
        max_r = max(r for r, c in all_cells.keys())
        max_c = max(c for r, c in all_cells.keys())

        # 初始化一个足够大的二维矩阵
        excel_data = [['' for _ in range(max_c + 1)] for _ in range(max_r + 1)]

        # 将字典中的内容填入矩阵
        for (r, c), text in all_cells.items():
            excel_data[r][c] = text

        # 导出为 Excel
        df = pd.DataFrame(excel_data)
        df = df.drop(df.columns[0], axis=1)
        try:
            # 1. 提取数据
            data_list = extract_data_from_header_footer_tables(word_file_path)

            if not data_list:
                print("警告：未在页眉/页脚的表格中找到信息。")
                # 为了演示，如果没找到，我们手动构造一行你提供的测试数据

            new_row_data = ["设备名称"]+data_list[8:9]+["设备类型"]+data_list[11:43]
            # 1. 将切片后的列表数据转换为一个单行的 DataFrame，并指定列名与原表一致
            new_row_df = pd.DataFrame([new_row_data], columns=df.columns)

            # 2. 使用 concat 将新行放在原 DataFrame 前面进行拼接
            # ignore_index=True 会重置索引，保证合并后索引是连续的 (0, 1, 2...)
            df = pd.concat([new_row_df, df], ignore_index=True)

            new_row_data =data_list[44]+"："+data_list[45]+"，"+data_list[46]+"："+data_list[47]
            df.at[len(df), df.columns[0]] = new_row_data

            #df.to_excel("output.xlsx", index=False, header=False)
            output_file = target_name.replace('.docx', '.xlsx')

            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                # 1. 将 DataFrame 写入 Excel (index=False 不写入行号)
                df.to_excel(writer, index=False, header=False,sheet_name='Sheet1')

                # 2. 获取 openpyxl 的工作簿和工作表对象
                workbook = writer.book
                worksheet = writer.sheets['Sheet1']

                # 3. 计算最后一行的行号
                last_row_num = len(df)

                # 4. 定义要合并的列范围（例如合并 A 列到 C 列）
                start_col = 1  # A列对应数字 1
                end_col = 33  # C列对应数字 3

                # 5. 执行合并操作
                # openpyxl 的 merge_cells 支持直接传入行列数字
                worksheet.merge_cells(start_row=last_row_num, start_column=start_col,
                                      end_row=last_row_num, end_column=end_col)

                # 6. (可选) 设置合并后的单元格居中对齐
                # 合并后，单元格的值会保留在最左上角（即 A 列的那个单元格）
                target_cell = worksheet.cell(row=last_row_num, column=start_col)
                target_cell.alignment = Alignment(horizontal='center', vertical='center')

                for row_index in range(2,last_row_num,2):
                    worksheet.merge_cells(start_row=row_index, start_column=1,
                                          end_row=row_index+1, end_column=1)
                    worksheet.merge_cells(start_row=row_index, start_column=2,
                                          end_row=row_index + 1, end_column=2)
                    worksheet.merge_cells(start_row=row_index, start_column=3,
                                          end_row=row_index + 1, end_column=3)
                    worksheet.merge_cells(start_row=row_index, start_column=35,
                                          end_row=row_index + 1, end_column=35)

        except Exception as e:
            print(f"处理过程中出错: {e}")
            print("请检查文件路径是否正确，或者文件是否被其他程序占用。")

    else:
        print("未在文档中提取到有效表格数据。")

else:
    print(f"错误：找不到名为 '{target_name}' 的文件，请检查文件名是否输入正确！")
    input("请按回车键退出...")